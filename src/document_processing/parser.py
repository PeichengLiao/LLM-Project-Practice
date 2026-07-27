"""
文档解析器: 支持 PDF、Word、Markdown、TXT 等多种格式
"""
import re
from pathlib import Path
from typing import List, Optional
from dataclasses import dataclass, field


@dataclass
class ParsedDocument:
    """解析后的文档结构"""
    content: str
    metadata: dict = field(default_factory=dict)
    source_path: Optional[str] = None
    page_count: int = 0


class DocumentParser:
    """多格式文档解析器"""

    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.md', '.txt', '.html'}

    def parse(self, file_path: str) -> ParsedDocument:
        """根据文件后缀选择对应的解析方法"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_FORMATS:
            raise ValueError(f"不支持的文件格式: {suffix}，支持的格式: {self.SUPPORTED_FORMATS}")

        parser_method = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_docx,
            '.md': self._parse_markdown,
            '.txt': self._parse_text,
            '.html': self._parse_html,
        }[suffix]

        return parser_method(file_path)

    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """解析 PDF 文件 (使用 PyMuPDF)"""
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        full_text = []
        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                # 标注页码
                full_text.append(f"[第{page_num + 1}页]\n{text}")

        return ParsedDocument(
            content="\n\n".join(full_text),
            metadata={
                "format": "pdf",
                "title": Path(file_path).stem,
                "file_path": file_path,
            },
            source_path=file_path,
            page_count=len(doc),
        )

    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """解析 Word 文档"""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        full_text = []

        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                # 识别标题
                if para.style.name.startswith("Heading"):
                    level = para.style.name.split()[-1]
                    prefix = "#" * int(level)
                    full_text.append(f"{prefix} {para.text}")
                else:
                    full_text.append(para.text)

        # 提取表格
        for table_idx, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                table_text.append(row_text)
            full_text.append(f"\n[表格{table_idx + 1}]\n" + "\n".join(table_text))

        return ParsedDocument(
            content="\n\n".join(full_text),
            metadata={
                "format": "docx",
                "title": Path(file_path).stem,
                "file_path": file_path,
            },
            source_path=file_path,
        )

    def _parse_markdown(self, file_path: str) -> ParsedDocument:
        """解析 Markdown 文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return ParsedDocument(
            content=content,
            metadata={"format": "markdown", "title": Path(file_path).stem},
            source_path=file_path,
        )

    def _parse_text(self, file_path: str) -> ParsedDocument:
        """解析纯文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return ParsedDocument(
            content=content,
            metadata={"format": "text", "title": Path(file_path).stem},
            source_path=file_path,
        )

    def _parse_html(self, file_path: str) -> ParsedDocument:
        """解析 HTML 文件"""
        from bs4 import BeautifulSoup

        with open(file_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'lxml')
            # 移除 script 和 style 标签
            for tag in soup(['script', 'style', 'nav', 'footer']):
                tag.decompose()
            content = soup.get_text(separator='\n', strip=True)

        return ParsedDocument(
            content=content,
            metadata={"format": "html", "title": Path(file_path).stem},
            source_path=file_path,
        )
